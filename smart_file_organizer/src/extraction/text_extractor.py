"""
Text Extraction Service
=======================

Extracts text content from various document formats including PDF, Word, and plain text.
Uses PyMuPDF for PDFs and python-docx for Word documents.
"""

from pathlib import Path
from typing import Optional, Tuple, List
from abc import ABC, abstractmethod
from dataclasses import dataclass
import mimetypes

from src.utils.logging_config import get_logger
from src.utils.exceptions import ExtractionError

logger = get_logger(__name__)

# Lazy imports for heavy libraries
fitz = None
Document = None


def _import_fitz():
    """Lazy import PyMuPDF."""
    global fitz
    if fitz is None:
        import fitz as _fitz
        fitz = _fitz
    return fitz


def _import_docx():
    """Lazy import python-docx."""
    global Document
    if Document is None:
        from docx import Document as _Document
        Document = _Document
    return Document


@dataclass
class ExtractionResult:
    """Result of text extraction.
    
    Attributes:
        text: Extracted text content.
        metadata: Document metadata.
        page_count: Number of pages (for multi-page documents).
        char_count: Number of characters extracted.
        extraction_method: Method used for extraction.
    """
    text: str
    metadata: dict
    page_count: int = 1
    char_count: int = 0
    extraction_method: str = "unknown"
    
    def __post_init__(self):
        """Calculate char count."""
        if self.char_count == 0:
            self.char_count = len(self.text)


class BaseExtractor(ABC):
    """Abstract base class for text extractors."""
    
    @property
    @abstractmethod
    def supported_extensions(self) -> set:
        """Set of supported file extensions."""
        pass
    
    @abstractmethod
    def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from file.
        
        Args:
            file_path: Path to the file.
        
        Returns:
            ExtractionResult with text and metadata.
        
        Raises:
            ExtractionError: If extraction fails.
        """
        pass
    
    def supports(self, file_path: Path) -> bool:
        """Check if this extractor supports the file.
        
        Args:
            file_path: Path to check.
        
        Returns:
            True if file is supported.
        """
        return file_path.suffix.lower() in self.supported_extensions


class PDFExtractor(BaseExtractor):
    """Extracts text from PDF files using PyMuPDF.
    
    Features:
    - Handles multi-page PDFs
    - Extracts document metadata
    - Early termination for large documents
    - Detection of scanned/image-only PDFs
    """
    
    def __init__(self, max_pages: int = 5, max_chars: int = 10000):
        """Initialize PDF extractor.
        
        Args:
            max_pages: Maximum pages to extract from.
            max_chars: Maximum characters to extract.
        """
        self.max_pages = max_pages
        self.max_chars = max_chars
    
    @property
    def supported_extensions(self) -> set:
        return {'.pdf'}
    
    def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file.
        
        Returns:
            ExtractionResult with extracted text.
        
        Raises:
            ExtractionError: If PDF cannot be processed.
        """
        fitz = _import_fitz()
        
        try:
            doc = fitz.open(file_path)
        except Exception as e:
            raise ExtractionError(
                f"Failed to open PDF: {e}",
                file_path=str(file_path),
                extractor_type="PDF"
            )
        
        try:
            text_parts = []
            total_chars = 0
            pages_with_text = 0
            
            metadata = {
                "title": doc.metadata.get("title", "") or "",
                "author": doc.metadata.get("author", "") or "",
                "subject": doc.metadata.get("subject", "") or "",
                "creator": doc.metadata.get("creator", "") or "",
                "producer": doc.metadata.get("producer", "") or "",
                "creation_date": doc.metadata.get("creationDate", "") or "",
            }
            
            for page_num in range(min(doc.page_count, self.max_pages)):
                page = doc[page_num]
                text = page.get_text("text")
                
                if text.strip():
                    text_parts.append(text)
                    total_chars += len(text)
                    pages_with_text += 1
                
                # Early termination if we have enough text
                if total_chars >= self.max_chars:
                    break
            
            combined_text = "\n".join(text_parts)
            
            # Check if likely scanned (no text extracted)
            is_scanned = pages_with_text == 0 and doc.page_count > 0
            if is_scanned:
                metadata["likely_scanned"] = True
                logger.info(f"PDF appears to be scanned/image-only: {file_path}")
            
            return ExtractionResult(
                text=combined_text,
                metadata=metadata,
                page_count=doc.page_count,
                extraction_method="PyMuPDF"
            )
            
        finally:
            doc.close()
    
    def is_scanned_pdf(self, file_path: Path) -> bool:
        """Check if PDF is likely scanned (no extractable text).
        
        Args:
            file_path: Path to PDF file.
        
        Returns:
            True if PDF appears to be scanned.
        """
        fitz = _import_fitz()
        try:
            doc = fitz.open(file_path)
            # Check first 3 pages
            for i in range(min(3, doc.page_count)):
                if doc[i].get_text("text").strip():
                    doc.close()
                    return False
            doc.close()
            return True
        except:
            return False


class WordExtractor(BaseExtractor):
    """Extracts text from Word documents using python-docx.
    
    Supports .docx format (not legacy .doc).
    """
    
    @property
    def supported_extensions(self) -> set:
        return {'.docx'}
    
    def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from Word document.
        
        Args:
            file_path: Path to Word document.
        
        Returns:
            ExtractionResult with extracted text.
        
        Raises:
            ExtractionError: If document cannot be processed.
        """
        Document = _import_docx()
        
        try:
            doc = Document(file_path)
        except Exception as e:
            raise ExtractionError(
                f"Failed to open Word document: {e}",
                file_path=str(file_path),
                extractor_type="Word"
            )
        
        # Extract paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_texts.append(" | ".join(row_text))
        
        # Combine all text
        all_text = paragraphs + table_texts
        combined_text = "\n".join(all_text)
        
        metadata = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "subject": doc.core_properties.subject or "",
            "keywords": doc.core_properties.keywords or "",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }
        
        return ExtractionResult(
            text=combined_text,
            metadata=metadata,
            extraction_method="python-docx"
        )


class PlainTextExtractor(BaseExtractor):
    """Extracts text from plain text files.
    
    Supports multiple text-based formats including TXT, MD, RST, etc.
    """
    
    SUPPORTED = {'.txt', '.md', '.markdown', '.rst', '.text', '.log'}
    
    def __init__(self, max_size: int = 1024 * 1024):  # 1MB default
        """Initialize text extractor.
        
        Args:
            max_size: Maximum file size to read in bytes.
        """
        self.max_size = max_size
    
    @property
    def supported_extensions(self) -> set:
        return self.SUPPORTED
    
    def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from plain text file.
        
        Args:
            file_path: Path to text file.
        
        Returns:
            ExtractionResult with file content.
        """
        file_size = file_path.stat().st_size
        
        # Detect encoding
        encoding = self._detect_encoding(file_path)
        
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                if file_size > self.max_size:
                    text = f.read(self.max_size)
                else:
                    text = f.read()
        except UnicodeDecodeError:
            # Fallback to latin-1 which accepts any byte sequence
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read(self.max_size)
        
        return ExtractionResult(
            text=text,
            metadata={
                "encoding": encoding,
                "original_size": file_size,
                "truncated": file_size > self.max_size,
            },
            extraction_method="plaintext"
        )
    
    def _detect_encoding(self, file_path: Path) -> str:
        """Detect file encoding.
        
        Args:
            file_path: Path to file.
        
        Returns:
            Detected encoding name.
        """
        # Read first chunk and try to detect
        try:
            with open(file_path, 'rb') as f:
                raw = f.read(4096)
            
            # Check for BOM
            if raw.startswith(b'\xef\xbb\xbf'):
                return 'utf-8-sig'
            if raw.startswith(b'\xff\xfe'):
                return 'utf-16-le'
            if raw.startswith(b'\xfe\xff'):
                return 'utf-16-be'
            
            # Try UTF-8
            try:
                raw.decode('utf-8')
                return 'utf-8'
            except UnicodeDecodeError:
                pass
            
            # Default to latin-1
            return 'latin-1'
            
        except Exception:
            return 'utf-8'


class TextExtractionService:
    """Main text extraction service.
    
    Coordinates multiple extractors to handle different file types.
    Falls back to OCR for scanned documents when needed.
    """
    
    def __init__(
        self,
        max_pages: int = 5,
        max_chars: int = 10000,
        ocr_engine = None
    ):
        """Initialize the extraction service.
        
        Args:
            max_pages: Maximum pages to extract from PDFs.
            max_chars: Maximum characters to extract.
            ocr_engine: Optional OCR engine for scanned documents.
        """
        self.extractors: List[BaseExtractor] = [
            PDFExtractor(max_pages=max_pages, max_chars=max_chars),
            WordExtractor(),
            PlainTextExtractor(),
        ]
        self.ocr_engine = ocr_engine
    
    def extract(self, file_path: Path) -> Optional[ExtractionResult]:
        """Extract text from a file.
        
        Args:
            file_path: Path to the file.
        
        Returns:
            ExtractionResult if successful, None if unsupported.
        
        Raises:
            ExtractionError: If extraction fails.
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise ExtractionError(
                "File does not exist",
                file_path=str(file_path)
            )
        
        for extractor in self.extractors:
            if extractor.supports(file_path):
                logger.debug(f"Using {type(extractor).__name__} for {file_path}")
                result = extractor.extract(file_path)
                
                # Check if we should fall back to OCR
                if (
                    isinstance(extractor, PDFExtractor) and
                    result.metadata.get("likely_scanned") and
                    self.ocr_engine is not None
                ):
                    logger.info(f"Falling back to OCR for scanned PDF: {file_path}")
                    return self._extract_with_ocr(file_path)
                
                return result
        
        logger.debug(f"No extractor supports file: {file_path}")
        return None
    
    def _extract_with_ocr(self, file_path: Path) -> ExtractionResult:
        """Extract text using OCR.
        
        Args:
            file_path: Path to the file.
        
        Returns:
            ExtractionResult from OCR.
        """
        # This will be implemented when OCR engine is integrated
        if self.ocr_engine:
            text = self.ocr_engine.extract_from_pdf(file_path)
            return ExtractionResult(
                text=text,
                metadata={"ocr_applied": True},
                extraction_method="OCR"
            )
        return ExtractionResult(text="", metadata={}, extraction_method="none")
    
    def supports(self, file_path: Path) -> bool:
        """Check if any extractor supports the file.
        
        Args:
            file_path: Path to check.
        
        Returns:
            True if file is supported.
        """
        file_path = Path(file_path)
        return any(e.supports(file_path) for e in self.extractors)
